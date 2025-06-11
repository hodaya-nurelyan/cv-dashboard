# backend/routes/integrations.py
from fastapi import APIRouter, Request, Response, Header
from fastapi.responses import StreamingResponse, JSONResponse, PlainTextResponse
from typing import Optional
import io
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import boto3
import os
import json

router = APIRouter()

s3_client = boto3.client('s3')
BUCKET_NAME = os.getenv('BUCKET_NAME', 'hodaya-cv-files')

CV_DATA = {
    "personal_info": {
        "name": "Hodaya Nurelyan",
        "title": "Frontend Developer",
        "email": "hodaya@example.com",
        "phone": "+972-XX-XXX-XXXX",
        "location": "Petah Tikva, Israel",
        "linkedin": "linkedin.com/in/hodaya-nurelyan",
        "github": "github.com/hodaya-n"
    },
    "summary": "Passionate Frontend Developer with expertise in React, Python, and AWS...",
    "experience": [
        {
            "company": "Tech Company",
            "position": "Frontend Developer",
            "duration": "2022-Present",
            "achievements": ["Built responsive web applications", "Improved performance by 30%"]
        }
    ],
    "skills": {
        "frontend": ["React", "JavaScript", "HTML/CSS", "TypeScript"],
        "backend": ["Python", "Flask", "Node.js"],
        "cloud": ["AWS", "Docker", "CI/CD"],
        "tools": ["Git", "VS Code", "Figma"]
    },
    "projects": [
        {
            "name": "Interactive CV Website",
            "description": "Built an AI-powered CV chat system",
            "technologies": ["React", "Python", "AWS"],
            "url": "https://askhodaya.com"
        }
    ]
}

def generate_pdf_cv():
    """Generate PDF version of CV"""
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Header
    p.setFont("Helvetica-Bold", 20)
    p.drawString(50, height - 50, CV_DATA["personal_info"]["name"])
    
    p.setFont("Helvetica", 14)
    p.drawString(50, height - 75, CV_DATA["personal_info"]["title"])
    
    # Contact Info
    y = height - 100
    p.setFont("Helvetica", 10)
    for key, value in CV_DATA["personal_info"].items():
        if key not in ["name", "title"]:
            p.drawString(50, y, f"{key.title()}: {value}")
            y -= 15
    
    # Summary
    y -= 20
    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, y, "Summary")
    y -= 20
    p.setFont("Helvetica", 10)
    
    # Split summary into multiple lines
    summary_lines = [CV_DATA["summary"][i:i+80] for i in range(0, len(CV_DATA["summary"]), 80)]
    for line in summary_lines:
        p.drawString(50, y, line)
        y -= 15
    
    # Experience
    y -= 20
    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, y, "Experience")
    y -= 20
    
    for exp in CV_DATA["experience"]:
        p.setFont("Helvetica-Bold", 10)
        p.drawString(50, y, f"{exp['position']} at {exp['company']}")
        y -= 15
        p.setFont("Helvetica", 10)
        p.drawString(50, y, exp["duration"])
        y -= 15
        for achievement in exp["achievements"]:
            p.drawString(70, y, f"• {achievement}")
            y -= 15
        y -= 10
    
    # Skills
    y -= 10
    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, y, "Skills")
    y -= 20
    
    for category, skills in CV_DATA["skills"].items():
        p.setFont("Helvetica-Bold", 10)
        p.drawString(50, y, f"{category.title()}:")
        p.setFont("Helvetica", 10)
        p.drawString(150, y, ", ".join(skills))
        y -= 15
    
    p.save()
    buffer.seek(0)
    return buffer


def generate_docx_cv():
    """Generate Word document version of CV"""
    doc = Document()
    
    # Header
    header = doc.add_heading(CV_DATA["personal_info"]["name"], 0)
    header.alignment = 1  # Center alignment
    
    subtitle = doc.add_heading(CV_DATA["personal_info"]["title"], level=2)
    subtitle.alignment = 1
    
    # Contact Information
    contact_info = doc.add_paragraph()
    for key, value in CV_DATA["personal_info"].items():
        if key not in ["name", "title"]:
            contact_info.add_run(f"{key.title()}: {value}\n")
    
    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(CV_DATA["summary"])
    
    # Experience
    doc.add_heading('Experience', level=1)
    for exp in CV_DATA["experience"]:
        exp_header = doc.add_paragraph()
        exp_header.add_run(f"{exp['position']} at {exp['company']}").bold = True
        doc.add_paragraph(exp["duration"])
        for achievement in exp["achievements"]:
            doc.add_paragraph(f"• {achievement}")
    
    # Skills
    doc.add_heading('Skills', level=1)
    for category, skills in CV_DATA["skills"].items():
        skill_para = doc.add_paragraph()
        skill_para.add_run(f"{category.title()}: ").bold = True
        skill_para.add_run(", ".join(skills))
    
    # Projects
    doc.add_heading('Projects', level=1)
    for project in CV_DATA["projects"]:
        proj_para = doc.add_paragraph()
        proj_para.add_run(project["name"]).bold = True
        doc.add_paragraph(project["description"])
        doc.add_paragraph(f"Technologies: {', '.join(project['technologies'])}")
        if project.get("url"):
            doc.add_paragraph(f"URL: {project['url']}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@router.get("/cv/download")
async def download_cv(format: str = "pdf"):
    try:
        format = format.lower()
        if format == "pdf":
            buffer = generate_pdf_cv()
            return StreamingResponse(buffer, media_type="application/pdf", headers={
                "Content-Disposition": "attachment; filename=Hodaya_Nurelyan_CV.pdf"
            })

        elif format == "docx":
            buffer = generate_docx_cv()
            return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
                "Content-Disposition": "attachment; filename=Hodaya_Nurelyan_CV.docx"
            })

        elif format == "json":
            return JSONResponse(CV_DATA)

        elif format == "txt":
            # יצירת גרסת טקסט פשוט
            text_content = f"{CV_DATA['personal_info']['name']}\n"
            # ... המשך יצירת טקסט ...
            return PlainTextResponse(text_content, headers={
                "Content-Disposition": "attachment; filename=Hodaya_Nurelyan_CV.txt"
            })

        return JSONResponse({"error": "Unsupported format. Use: pdf, docx, json, txt"}, status_code=400)
    except Exception as e:
        print(f"Error generating CV: {e}")
        return JSONResponse({"error": "Failed to generate CV"}, status_code=500)

@router.post("/analytics/download")
async def track_download(request: Request):
    try:
        data = await request.json()
        analytics_data = {
            'timestamp': datetime.now().isoformat(),
            'format': data.get('format', 'unknown'),
            'user_ip': request.client.host,
            'user_agent': request.headers.get('user-agent', ''),
            'referrer': request.headers.get('referer', ''),
            'session_id': data.get('session_id')
        }
        print(f"Download tracked: {analytics_data}")
        return JSONResponse({"status": "tracked", "timestamp": analytics_data["timestamp"]})
    except Exception as e:
        print(f"Error tracking download: {e}")
        return JSONResponse({"error": "Failed to track download"}, status_code=500)

@router.post("/scheduling/webhook")
async def calendly_webhook(request: Request):
    try:
        event_data = await request.json()
        if event_data.get('event') == 'invitee.created':
            invitee = event_data.get('payload', {}).get('invitee', {})
            meeting_data = {
                'timestamp': datetime.now().isoformat(),
                'invitee_email': invitee.get('email'),
                'invitee_name': invitee.get('name'),
                'event_type': event_data.get('payload', {}).get('event_type', {}).get('name'),
                'scheduled_time': invitee.get('scheduled_event', {}).get('start_time'),
                'meeting_url': invitee.get('scheduled_event', {}).get('location', {}).get('join_url')
            }
            print(f"New meeting scheduled: {meeting_data}")
        return JSONResponse({"status": "processed"})
    except Exception as e:
        print(f"Error processing webhook: {e}")
        return JSONResponse({"error": "Failed to process webhook"}, status_code=500)

@router.post("/share/track")
async def track_share(request: Request):
    try:
        data = await request.json()
        share_data = {
            'timestamp': datetime.now().isoformat(),
            'platform': data.get('platform', 'unknown'),
            'user_ip': request.client.host,
            'referrer': request.headers.get('referer', ''),
            'user_agent': request.headers.get('user-agent', '')
        }
        print(f"Share tracked: {share_data}")
        return JSONResponse({"status": "tracked"})
    except Exception as e:
        print(f"Error tracking share: {e}")
        return JSONResponse({"error": "Failed to track share"}, status_code=500)

@router.post("/contact/quick")
async def quick_contact(request: Request):
    try:
        data = await request.json()
        if not data or not data.get("email"):
            return JSONResponse({"error": "Email is required"}, status_code=400)
        contact_data = {
            'timestamp': datetime.now().isoformat(),
            'name': data.get('name', ''),
            'email': data.get('email'),
            'message': data.get('message', ''),
            'source': 'quick_contact',
            'user_ip': request.client.host
        }
        print(f"New contact: {contact_data}")
        return JSONResponse({"status": "message_sent", "timestamp": contact_data["timestamp"]})
    except Exception as e:
        print(f"Error processing contact: {e}")
        return JSONResponse({"error": "Failed to send message"}, status_code=500)

@router.get("/health")
async def health_check():
    return JSONResponse({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0"
    })
